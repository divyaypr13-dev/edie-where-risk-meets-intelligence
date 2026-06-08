"""
EDIE Enterprise - Complete Production Backend
Contract-Type Specific Logic for Realistic Demo
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from contextlib import asynccontextmanager
from datetime import datetime
import os
import re
import json
import uuid
import logging
from io import BytesIO
from typing import List, Optional, Dict, Any
import requests

# ============================================
# CONFIGURATION
# ============================================

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
AI_AVAILABLE = bool(OPENROUTER_API_KEY)

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================
# AI FUNCTIONS
# ============================================

def call_ai(prompt: str, max_tokens: int = 800, temperature: float = 0.7, 
            system_prompt: str = None, conversation_history: List = None) -> Optional[str]:
    if not OPENROUTER_API_KEY:
        return None
    
    try:
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        else:
            messages.append({"role": "system", "content": "You are EDIE, a friendly, knowledgeable AI assistant specializing in Indian laws, contract clauses, and compliance. Be conversational and helpful."})
        
        if conversation_history:
            for msg in conversation_history[-8:]:
                messages.append(msg)
        
        messages.append({"role": "user", "content": prompt})
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "google/gemma-2-9b-it:free",
                "messages": messages,
                "max_tokens": max_tokens,
                "temperature": temperature
            },
            timeout=60
        )
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        return None
    except Exception as e:
        logger.error(f"AI Error: {e}")
        return None

def call_ai_analysis(question: str, analysis_data: Dict, conversation_history: List = None) -> Optional[str]:
    if not OPENROUTER_API_KEY:
        return None
    
    risk = analysis_data.get('risk_assessment', {})
    clauses = analysis_data.get('clause_analysis', {})
    fin = analysis_data.get('financial_analysis', {})
    compliance = analysis_data.get('compliance', {})
    
    system_prompt = f"""You are EDIE, a contract analysis expert. Explain THIS specific contract:

RISK SCORE: {risk.get('score', 0)}/100 ({risk.get('level', 'Unknown')})
CONTRACT TYPE: {analysis_data.get('contract_type', 'General')}
DECISION: {risk.get('decision', 'Unknown')}
MISSING CLAUSES: {', '.join(clauses.get('missing_list', [])[:8])}
PRESENT CLAUSES: {', '.join(clauses.get('present_list', [])[:5])}
RECOMMENDATIONS: {', '.join(analysis_data.get('recommendations', [])[:5])}
ROI: {fin.get('roi_percentage', 0)}%

Answer questions conversationally about THIS specific contract analysis. Be helpful and specific."""

    try:
        messages = [{"role": "system", "content": system_prompt}]
        if conversation_history:
            for msg in conversation_history[-10:]:
                messages.append(msg)
        messages.append({"role": "user", "content": question})
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"},
            json={"model": "google/gemma-2-9b-it:free", "messages": messages, "max_tokens": 600, "temperature": 0.7},
            timeout=60
        )
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        return None
    except Exception as e:
        logger.error(f"AI Analysis Error: {e}")
        return None

# ============================================
# KNOWLEDGE BASE LOADER
# ============================================

def clean_json_content(content: str) -> str:
    lines = content.split('\n')
    cleaned_lines = []
    in_string = False
    for line in lines:
        i = 0
        new_line = []
        while i < len(line):
            if line[i] == '"' and (i == 0 or line[i-1] != '\\'):
                in_string = not in_string
                new_line.append(line[i])
            elif not in_string and line[i:i+2] == '//':
                break
            else:
                new_line.append(line[i])
            i += 1
        cleaned_lines.append(''.join(new_line))
    content = '\n'.join(cleaned_lines)
    content = re.sub(r'/\*.*?\*/', '', content, flags=re.DOTALL)
    return content

def safe_json_load(filepath: str) -> Optional[Dict]:
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        cleaned = clean_json_content(content)
        return json.loads(cleaned)
    except Exception as e:
        logger.warning(f"Error loading {filepath}: {e}")
        return None

class KnowledgeBase:
    def __init__(self, kb_path="Knowledge_base"):
        self.kb_path = kb_path
        self.loaded_files = []
        self._load_all_files()
    
    def _load_all_files(self):
        if not os.path.exists(self.kb_path):
            logger.warning(f"Knowledge base folder not found: {self.kb_path}")
            return
        for filename in os.listdir(self.kb_path):
            if not filename.endswith('.json'):
                continue
            filepath = os.path.join(self.kb_path, filename)
            data = safe_json_load(filepath)
            if data:
                self.loaded_files.append(filename)
                logger.info(f"Loaded: {filename}")
        logger.info(f"Knowledge base loaded: {len(self.loaded_files)} files")

kb = KnowledgeBase()

# ============================================
# STORAGE
# ============================================

class Storage:
    def __init__(self):
        self.analyses = []
        self.chat_sessions = {}
        self.analysis_chat_sessions = {}
    
    def add_analysis(self, data):
        data['id'] = str(uuid.uuid4())
        data['created_at'] = datetime.now().isoformat()
        self.analyses.insert(0, data)
        if len(self.analyses) > 50:
            self.analyses = self.analyses[:50]
        return data
    
    def get_all(self, limit=20):
        return self.analyses[:limit]

storage = Storage()

# ============================================
# IMPROVED CONTRACT TYPE DETECTION
# ============================================

def detect_contract_type(text: str) -> str:
    text_lower = text.lower()
    
    # Strong pattern matching with multiple keywords
    detection_rules = [
        {"type": "Employment Agreement", "keywords": ["employment agreement", "employee", "employer", "salary", "compensation", "probation", "notice period", "job title", "annual salary", "employee shall"], "weight": 3},
        {"type": "NDA", "keywords": ["non-disclosure", "nda", "confidentiality agreement", "confidential information", "trade secret", "disclosing party", "receiving party"], "weight": 3},
        {"type": "Software Agreement", "keywords": ["software", "license", "source code", "software agreement", "licensor", "licensee", "software license", "intellectual property rights"], "weight": 3},
        {"type": "Service Agreement", "keywords": ["service agreement", "services", "service provider", "sla", "service level", "response time", "uptime"], "weight": 2},
        {"type": "Consulting Agreement", "keywords": ["consulting", "consultant", "consulting services", "independent contractor", "professional services"], "weight": 2},
        {"type": "Vendor Agreement", "keywords": ["vendor", "supplier", "vendor agreement", "supply agreement", "procurement", "purchase order"], "weight": 2},
        {"type": "Purchase Agreement", "keywords": ["purchase", "buyer", "seller", "goods", "products", "delivery", "price", "shipment"], "weight": 2},
        {"type": "Lease Agreement", "keywords": ["lease", "rent", "lessor", "lessee", "premises", "security deposit", "rental"], "weight": 3},
        {"type": "Maintenance Agreement", "keywords": ["maintenance", "amc", "annual maintenance", "service provider", "breakdown", "repair"], "weight": 2},
        {"type": "Distribution Agreement", "keywords": ["distribution", "distributor", "exclusive distribution", "territory", "distribute"], "weight": 2},
        {"type": "Joint Venture Agreement", "keywords": ["joint venture", "venture", "partnership", "jointly", "collaboration agreement"], "weight": 2},
    ]
    
    scores = {}
    for rule in detection_rules:
        score = 0
        for kw in rule["keywords"]:
            if kw in text_lower:
                score += rule["weight"]
        if score > 0:
            scores[rule["type"]] = score
    
    if scores:
        # Return the highest scoring contract type
        detected_type = max(scores, key=scores.get)
        logger.info(f"Detected contract type: {detected_type} (score: {scores[detected_type]})")
        return detected_type
    
    # Check for specific document titles
    lines = text_lower.split('\n')
    for line in lines[:5]:
        if "agreement" in line:
            if "employment" in line:
                return "Employment Agreement"
            elif "confidential" in line or "non-disclosure" in line:
                return "NDA"
            elif "software" in line or "license" in line:
                return "Software Agreement"
            elif "service" in line:
                return "Service Agreement"
            elif "consulting" in line:
                return "Consulting Agreement"
            elif "vendor" in line or "supplier" in line:
                return "Vendor Agreement"
            elif "purchase" in line:
                return "Purchase Agreement"
            elif "lease" in line:
                return "Lease Agreement"
            elif "maintenance" in line:
                return "Maintenance Agreement"
            elif "distribution" in line:
                return "Distribution Agreement"
            elif "joint" in line:
                return "Joint Venture Agreement"
    
    return "General Commercial Agreement"

# ============================================
# CONTRACT TYPE → MANDATORY CHECKS MATRIX
# ============================================

CONTRACT_TYPE_CONFIG = {
    "Consulting Agreement": {
        "mandatory_clauses": ["Parties", "Scope of Work", "Payment Terms", "Termination", "Liability", "Confidentiality", "Intellectual Property"],
        "laws": ["GST", "TDS", "MSME"],
        "cyber_privacy": ["DPDP"]
    },
    "Distribution Agreement": {
        "mandatory_clauses": ["Parties", "Payment Terms", "Termination", "Liability", "Governing Law"],
        "laws": ["GST", "MSME"],
        "cyber_privacy": []
    },
    "Employment Agreement": {
        "mandatory_clauses": ["Parties", "Scope of Work", "Compensation", "Termination", "Confidentiality", "Intellectual Property"],
        "laws": ["Labour Codes", "Income Tax", "DPDP"],
        "cyber_privacy": ["DPDP"]
    },
    "Joint Venture Agreement": {
        "mandatory_clauses": ["Parties", "Ownership", "Profit Sharing", "Liability", "Dispute Resolution", "Termination"],
        "laws": ["Companies Act", "GST"],
        "cyber_privacy": []
    },
    "Lease Agreement": {
        "mandatory_clauses": ["Parties", "Rent", "Security Deposit", "Term", "Maintenance", "Termination"],
        "laws": ["Contract Act"],
        "cyber_privacy": []
    },
    "Maintenance Agreement": {
        "mandatory_clauses": ["Parties", "Scope", "SLA", "Payment Terms", "Liability", "Termination"],
        "laws": ["GST"],
        "cyber_privacy": ["DPDP"]
    },
    "NDA": {
        "mandatory_clauses": ["Parties", "Definition", "Confidentiality", "Term", "Exceptions", "Remedies"],
        "laws": ["Contract Act"],
        "cyber_privacy": ["DPDP"]
    },
    "Purchase Agreement": {
        "mandatory_clauses": ["Parties", "Goods", "Price", "Delivery", "Warranty", "Liability", "Payment Terms"],
        "laws": ["GST", "MSME"],
        "cyber_privacy": []
    },
    "Service Agreement": {
        "mandatory_clauses": ["Parties", "Scope", "SLA", "Payment Terms", "Liability", "Termination", "Confidentiality"],
        "laws": ["GST", "TDS", "MSME"],
        "cyber_privacy": ["DPDP"]
    },
    "Software Agreement": {
        "mandatory_clauses": ["Parties", "License", "Intellectual Property", "Payment Terms", "Liability", "SLA", "Confidentiality", "Data Protection"],
        "laws": ["GST", "DPDP", "IT Act"],
        "cyber_privacy": ["DPDP", "ISO 27001"]
    },
    "Vendor Agreement": {
        "mandatory_clauses": ["Parties", "Scope", "Payment Terms", "Liability", "Indemnity", "Termination"],
        "laws": ["GST", "MSME"],
        "cyber_privacy": ["DPDP"]
    },
    "General Commercial Agreement": {
        "mandatory_clauses": ["Parties", "Scope", "Payment Terms", "Termination", "Liability", "Governing Law"],
        "laws": ["Contract Act", "GST"],
        "cyber_privacy": []
    }
}

# ============================================
# STANDARD CLAUSE KEYWORDS
# ============================================

CLAUSE_KEYWORDS = {
    "Parties": ["party", "company", "corporation", "between", "employer", "employee", "hereinafter"],
    "Scope of Work": ["scope", "services", "deliverable", "work", "provide", "duties", "responsibilities", "obligations"],
    "Payment Terms": ["payment", "fee", "price", "invoice", "pay", "compensation", "salary", "remuneration"],
    "Term": ["term", "duration", "commencement", "effective date", "period", "validity"],
    "Termination": ["termination", "notice period", "resignation", "dismissal", "end", "expiry", "cure period"],
    "Liability": ["liability", "cap", "damage", "claim", "loss", "limitation", "maximum liability"],
    "Governing Law": ["governing law", "laws of", "choice of law", "governed by"],
    "Jurisdiction": ["jurisdiction", "courts", "venue", "exclusive jurisdiction", "competent courts"],
    "Dispute Resolution": ["dispute", "arbitration", "mediation", "litigation", "arbitrator"],
    "Indemnity": ["indemnity", "indemnify", "hold harmless", "defend"],
    "Confidentiality": ["confidential", "nda", "non-disclosure", "secret", "trade secret", "proprietary"],
    "Force Majeure": ["force majeure", "act of god", "unforeseeable", "emergency", "beyond control"],
    "Intellectual Property": ["intellectual property", "ip", "ownership", "copyright", "patent", "trademark", "invention"],
    "Warranty": ["warranty", "guarantee", "represent", "undertake", "merchantability"],
    "Audit Rights": ["audit", "inspect", "records", "verify", "review", "examine"],
    "Insurance": ["insurance", "coverage", "liability insurance", "indemnify"],
    "SLA": ["sla", "service level", "uptime", "response time", "resolution time", "service credit"],
    "Compensation": ["salary", "compensation", "remuneration", "wage", "bonus", "incentive"],
    "Rent": ["rent", "lease amount", "monthly rent", "rental"],
    "Security Deposit": ["security deposit", "deposit", "advance", "earnest money"],
    "Maintenance": ["maintenance", "repair", "upkeep", "service", "amc"],
    "Data Protection": ["data protection", "privacy", "personal data", "dpdp", "gdpr", "data processing"],
    "Non-Compete": ["non-compete", "non competition", "competitive activity", "restrictive covenant"],
    "Non-Solicit": ["non-solicit", "solicit", "poach", "hire away"],
}

# ============================================
# LAW/COMPLIANCE KEYWORDS
# ============================================

LAW_KEYWORDS = {
    "GST": ["gst", "goods and services tax", "tax invoice", "input tax credit", "gstin", "cgst", "sgst", "igst"],
    "TDS": ["tds", "tax deducted at source", "withholding tax", "section 194", "deduct tax"],
    "MSME": ["msme", "micro small", "45 days", "delayed payment", "facilitation council", "udyam"],
    "DPDP": ["data protection", "personal data", "consent", "privacy", "dpdp", "data fiduciary", "data principal"],
    "IT Act": ["it act", "information technology act", "cyber", "electronic record", "digital signature"],
    "Labour Codes": ["labor", "labour", "employee rights", "workplace", "industrial dispute", "minimum wage", "gratuity", "provident fund"],
    "Income Tax": ["income tax", "tax deducted", "tan", "pan", "form 16", "form 26q"],
    "Companies Act": ["companies act", "board resolution", "director", "annual report", "roc"],
    "Contract Act": ["contract act", "indian contract", "void", "enforceable", "consideration"],
    "ISO 27001": ["information security", "security policy", "security measures", "access control", "incident management", "business continuity"],
}

# ============================================
# CLAUSE CHECKING FUNCTION
# ============================================

def check_clauses(text: str, contract_type: str) -> Dict:
    text_lower = text.lower()
    
    # Get mandatory clauses for this contract type
    config = CONTRACT_TYPE_CONFIG.get(contract_type, CONTRACT_TYPE_CONFIG["General Commercial Agreement"])
    mandatory_clauses = config.get("mandatory_clauses", [])
    
    present_clauses = []
    missing_clauses = []
    present_all = []
    missing_all = []
    
    # Check all standard clauses
    for clause_name, keywords in CLAUSE_KEYWORDS.items():
        found = any(kw in text_lower for kw in keywords)
        if found:
            present_all.append(clause_name)
            if clause_name in mandatory_clauses:
                present_clauses.append(clause_name)
        else:
            missing_all.append(clause_name)
            if clause_name in mandatory_clauses:
                missing_clauses.append(clause_name)
    
    # Calculate clause score (0-30 points based on mandatory clauses present)
    if mandatory_clauses:
        present_mandatory = sum(1 for c in mandatory_clauses if c in present_all)
        clause_score = (present_mandatory / len(mandatory_clauses)) * 30
    else:
        present_mandatory = 0
        clause_score = 15
    
    return {
        "present": present_clauses,
        "missing": missing_clauses,
        "present_all": present_all,
        "missing_all": missing_all,
        "score": clause_score,
        "mandatory_total": len(mandatory_clauses),
        "mandatory_present": present_mandatory
    }

# ============================================
# COMPLIANCE CHECKING FUNCTION
# ============================================

def check_compliance(text: str, contract_type: str) -> Dict:
    text_lower = text.lower()
    
    config = CONTRACT_TYPE_CONFIG.get(contract_type, CONTRACT_TYPE_CONFIG["General Commercial Agreement"])
    laws_to_check = config.get("laws", [])
    cyber_to_check = config.get("cyber_privacy", [])
    
    all_standards = laws_to_check + cyber_to_check
    
    if not all_standards:
        all_standards = ["Contract Act"]
    
    results = []
    total_score = 0
    
    for standard in all_standards:
        keywords = LAW_KEYWORDS.get(standard, [])
        if keywords:
            found = sum(1 for kw in keywords if kw in text_lower)
            score = round((found / len(keywords)) * 100)
            score = min(100, score)
        else:
            score = 50
        
        total_score += score
        status = "Compliant" if score >= 70 else "Partial" if score >= 40 else "Non-Compliant"
        results.append({"standard": standard, "status": status, "score": score})
    
    overall = round(total_score / len(all_standards)) if all_standards else 50
    
    return {"overall_score": overall, "standards": results}

# ============================================
# RISK PATTERN DETECTION
# ============================================

def detect_risks(text: str, contract_type: str) -> tuple:
    text_lower = text.lower()
    risk_contributors = []
    pattern_penalty = 0
    
    risk_patterns = [
        {"name": "Unlimited Liability", "keywords": ["unlimited liability", "without limitation", "any and all damages"], "increase": 20, "severity": "Critical", 
         "suggestion": "Add liability cap of 12 months fees or fixed amount (₹10-50 lakhs)"},
        {"name": "Unlimited Indemnity", "keywords": ["indemnify.*without limit", "unlimited indemnity"], "increase": 15, "severity": "Critical",
         "suggestion": "Cap indemnity at fees paid or fixed amount, make mutual"},
        {"name": "Missing Governing Law", "keywords": [], "increase": 12, "severity": "High", "detection_type": "absence",
         "suggestion": "Add governing law clause specifying Indian law"},
        {"name": "Missing Jurisdiction", "keywords": [], "increase": 10, "severity": "High", "detection_type": "absence",
         "suggestion": "Add jurisdiction clause specifying courts in India"},
        {"name": "Foreign Governing Law", "keywords": ["governed by laws of (?!india)", "laws of (?!india)"], "increase": 15, "severity": "Critical",
         "suggestion": "Change governing law to India"},
        {"name": "Foreign Jurisdiction", "keywords": ["jurisdiction.*(?!india)", "courts in (?!india)"], "increase": 12, "severity": "High",
         "suggestion": "Change jurisdiction to Indian courts"},
        {"name": "Missing Termination", "keywords": [], "increase": 10, "severity": "High", "detection_type": "absence",
         "suggestion": "Add termination clause with reasonable notice period"},
        {"name": "Missing Confidentiality", "keywords": [], "increase": 10, "severity": "High", "detection_type": "absence",
         "suggestion": "Add confidentiality clause"},
        {"name": "Missing Liability Cap", "keywords": [], "increase": 15, "severity": "Critical", "detection_type": "absence",
         "suggestion": "Add mutual liability cap clause"},
        {"name": "Missing Data Protection", "keywords": [], "increase": 15, "severity": "Critical", "detection_type": "absence",
         "suggestion": "Add DPDP Act compliant data protection clause"},
        {"name": "Auto Renewal", "keywords": ["auto renew", "automatically renew", "automatic renewal"], "increase": 8, "severity": "Medium",
         "suggestion": "Change notice period to 60-90 days or make renewal mutual"},
        {"name": "One-sided Termination", "keywords": ["terminate immediately", "without notice", "sole discretion"], "increase": 8, "severity": "Medium",
         "suggestion": "Make termination rights mutual with reasonable notice"},
        {"name": "AS IS Warranty", "keywords": ["as is", "without warranty", "no warranty", "with all faults"], "increase": 10, "severity": "High",
         "suggestion": "Add limited warranty for 30-90 days"},
        {"name": "No SLA Service Credits", "keywords": [], "increase": 8, "severity": "Medium", "detection_type": "absence",
         "suggestion": "Add service credits for SLA breaches"},
    ]
    
    for risk in risk_patterns:
        found = False
        keywords = risk.get("keywords", [])
        
        if keywords:
            for kw in keywords:
                try:
                    if re.search(kw, text_lower, re.IGNORECASE):
                        found = True
                        break
                except:
                    if kw in text_lower:
                        found = True
                        break
        elif risk.get("detection_type") == "absence":
            clause_name = risk["name"].replace("Missing ", "")
            if clause_name not in text_lower:
                found = True
        
        if found:
            risk_contributors.append({
                "pattern": risk["name"],
                "increase": risk["increase"],
                "severity": risk["severity"],
                "suggestion": risk["suggestion"]
            })
            pattern_penalty += risk["increase"]
    
    pattern_score = min(30, pattern_penalty)
    return risk_contributors, pattern_score

# ============================================
# MAIN ANALYSIS FUNCTION
# ============================================

def analyze_contract(text: str, contract_value: float = 10000000, term_years: int = 2, industry: str = "default"):
    
    # Step 1: Detect contract type
    contract_type = detect_contract_type(text)
    logger.info(f"Contract Type Detected: {contract_type}")
    
    # Step 2: Check clauses
    clause_result = check_clauses(text, contract_type)
    
    # Step 3: Check compliance (contract-type specific)
    compliance_result = check_compliance(text, contract_type)
    
    # Step 4: Detect risk patterns
    risk_contributors, pattern_score = detect_risks(text, contract_type)
    
    # Step 5: Calculate risk score
    clause_score = clause_result["score"]
    compliance_score = (100 - compliance_result["overall_score"]) * 0.20
    base_risk = clause_score + pattern_score + compliance_score
    
    industry_multipliers = {"banking": 1.3, "healthcare": 1.25, "government": 1.2, "technology": 1.1, "default": 1.0}
    multiplier = industry_multipliers.get(industry.lower(), 1.0)
    final_risk = min(100, base_risk * multiplier)
    
    # Step 6: Determine risk level and decision
    if final_risk >= 80:
        level = "CRITICAL"
        decision = "REJECT"
        action = "Do not sign. Critical risks require major renegotiation or rejection."
    elif final_risk >= 65:
        level = "HIGH"
        decision = "NEGOTIATE"
        action = "Major negotiations required. Do not sign without significant changes."
    elif final_risk >= 45:
        level = "MEDIUM"
        decision = "REVIEW"
        action = "Review carefully with legal team. Moderate risks present."
    elif final_risk >= 25:
        level = "LOW"
        decision = "LOW RISK"
        action = "Low risk. Can sign with standard due diligence."
    else:
        level = "MINIMAL"
        decision = "APPROVE"
        action = "Safe to sign. Well-drafted contract."
    
    # Step 7: Generate summary
    summary = f"This {contract_type} has a {level.lower()} risk score of {round(final_risk)}/100. "
    summary += f"It includes {clause_result['mandatory_present']} of {clause_result['mandatory_total']} mandatory clauses. "
    if clause_result["missing"]:
        summary += f"Missing clauses include {', '.join(clause_result['missing'][:3])}. "
    if risk_contributors:
        summary += f"Key risks: {', '.join([r['pattern'] for r in risk_contributors[:2]])}."
    
    # Step 8: Generate recommendations
    recommendations = []
    for clause in clause_result["missing"][:4]:
        recommendations.append(f"Add {clause} clause")
    for risk in risk_contributors[:3]:
        recommendations.append(risk["suggestion"])
    if not recommendations:
        recommendations = ["Review all terms before signing", "Standard due diligence recommended"]
    
    # Step 9: Financial Analysis
    total_revenue = contract_value * term_years
    expected_profit = total_revenue * 0.30
    missing_ratio = len(clause_result["missing_all"]) / max(len(clause_result["missing_all"]), 1)
    risk_exposure = missing_ratio * contract_value * 0.03 * (final_risk / 50)
    risk_adjusted_profit = max(0, expected_profit - risk_exposure)
    roi = (risk_adjusted_profit / total_revenue * 100) if total_revenue > 0 else 0
    
    legal_risk = min(100, final_risk + 5)
    financial_risk = min(100, final_risk + 8 if "Payment Terms" in clause_result["missing_all"] else final_risk)
    compliance_risk = min(100, final_risk + 12 if "Data Protection" in clause_result["missing_all"] else final_risk)
    operational_risk = max(0, final_risk - 10 if "SLA" in clause_result["present_all"] else final_risk)
    
    return {
        "risk_assessment": {
            "score": round(final_risk), "level": level, "decision": decision, "action": action,
            "legal_risk": round(legal_risk), "financial_risk": round(financial_risk),
            "compliance_risk": round(compliance_risk), "operational_risk": round(operational_risk)
        },
        "clause_analysis": {
            "total": len(clause_result["missing_all"]), "present": len(clause_result["present_all"]), 
            "missing": len(clause_result["missing_all"]),
            "present_list": clause_result["present_all"], "missing_list": clause_result["missing_all"],
            "mandatory_present": clause_result["mandatory_present"], "mandatory_total": clause_result["mandatory_total"]
        },
        "risk_contributors": risk_contributors,
        "recommendations": recommendations[:6],
        "summary": summary,
        "final_action": action,
        "contract_type": contract_type,
        "compliance": compliance_result,
        "financial_analysis": {
            "contract_value": contract_value, "term_years": term_years,
            "total_revenue": total_revenue, "expected_profit": round(expected_profit, 2),
            "risk_exposure": round(risk_exposure, 2), "risk_adjusted_profit": round(risk_adjusted_profit, 2),
            "roi_percentage": round(roi, 1)
        }
    }

# ============================================
# FASTAPI APP
# ============================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("EDIE Enterprise Starting...")
    logger.info(f"AI Available: {AI_AVAILABLE and bool(OPENROUTER_API_KEY)}")
    logger.info("Contract-Type Specific Analysis Enabled")
    yield

app = FastAPI(title="EDIE Enterprise", version="5.0.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

chat_sessions = {}
analysis_sessions = {}

@app.get("/")
async def root():
    return {"name": "EDIE Enterprise", "status": "running", "ai_available": AI_AVAILABLE and bool(OPENROUTER_API_KEY)}

@app.get("/dashboard")
async def get_dashboard():
    analyses = storage.get_all()
    total = len(analyses)
    high_risk = sum(1 for a in analyses if a.get('risk_assessment', {}).get('score', 0) >= 65)
    avg_risk = sum(a.get('risk_assessment', {}).get('score', 0) for a in analyses) / max(total, 1)
    return {
        "total_documents_analyzed": total, "high_risk_contracts": high_risk,
        "critical_contracts": sum(1 for a in analyses if a.get('risk_assessment', {}).get('score', 0) >= 80),
        "average_risk_score": round(avg_risk, 1),
        "risk_distribution": {
            "critical": sum(1 for a in analyses if a.get('risk_assessment', {}).get('score', 0) >= 80),
            "high": sum(1 for a in analyses if 65 <= a.get('risk_assessment', {}).get('score', 0) < 80),
            "medium": sum(1 for a in analyses if 45 <= a.get('risk_assessment', {}).get('score', 0) < 65),
            "low": sum(1 for a in analyses if 25 <= a.get('risk_assessment', {}).get('score', 0) < 45),
            "minimal": sum(1 for a in analyses if a.get('risk_assessment', {}).get('score', 0) < 25)
        },
        "recent_uploads": [{"id": a.get('id'), "filename": a.get('filename', 'Input'), "risk_score": a.get('risk_assessment', {}).get('score', 0), "date": a.get('created_at', '')[:10]} for a in analyses[:5]]
    }

@app.post("/analyze")
async def analyze_contract_endpoint(req: dict):
    try:
        result = analyze_contract(req.get('contract_text', ''), req.get('contract_value', 10000000), req.get('term_years', 2), req.get('industry', 'default'))
        
        if AI_AVAILABLE and OPENROUTER_API_KEY:
            ai_prompt = f"Risk {result['risk_assessment']['score']}/100 ({result['risk_assessment']['level']}). This is a {result.get('contract_type', 'general')} contract. Brief friendly summary (2 sentences)."
            ai_explanation = call_ai(ai_prompt, max_tokens=150, temperature=0.5)
            if ai_explanation:
                result["ai_explanation"] = ai_explanation
        
        stored = storage.add_analysis({"filename": "direct_input", "risk_assessment": result["risk_assessment"], "compliance": result["compliance"]})
        result["analysis_id"] = stored["id"]
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/analyze-file")
async def analyze_file(file: UploadFile = File(...), contract_value: float = Form(10000000), term_years: int = Form(2), industry: str = Form("default")):
    try:
        content = await file.read()
        ext = file.filename.split('.')[-1].lower()
        text = ""
        if ext == 'txt':
            text = content.decode('utf-8', errors='ignore')
        elif ext == 'pdf' and PYMUPDF_AVAILABLE:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
        elif ext == 'docx' and DOCX_AVAILABLE:
            doc = Document(BytesIO(content))
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            text = content.decode('utf-8', errors='ignore')[:5000]
        
        if len(text.strip()) < 50:
            return {"success": False, "error": "Could not extract enough text"}
        
        result = analyze_contract(text, contract_value, term_years, industry)
        
        if AI_AVAILABLE and OPENROUTER_API_KEY:
            ai_prompt = f"Risk {result['risk_assessment']['score']}/100 ({result['risk_assessment']['level']}). This is a {result.get('contract_type', 'general')} contract. Brief friendly summary."
            ai_explanation = call_ai(ai_prompt, max_tokens=150, temperature=0.5)
            if ai_explanation:
                result["ai_explanation"] = ai_explanation
        
        stored = storage.add_analysis({"filename": file.filename, "risk_assessment": result["risk_assessment"], "compliance": result["compliance"]})
        result["analysis_id"] = stored["id"]
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/ai/chat")
async def ai_chat(question: str = Form(...), session_id: str = Form(None)):
    if not OPENROUTER_API_KEY:
        return {"success": False, "error": "OpenRouter API key not configured."}
    
    if not session_id:
        session_id = str(uuid.uuid4())
    if session_id not in chat_sessions:
        chat_sessions[session_id] = []
    
    answer = call_ai(question, max_tokens=600, temperature=0.7, conversation_history=chat_sessions[session_id])
    
    if not answer:
        q_lower = question.lower()
        if 'dpdp' in q_lower:
            answer = "The DPDP Act 2023 is India's data protection law. Penalties up to ₹250 crore for violations. Requirements include consent, breach notification within 48 hours, and giving individuals rights to access, correct, and erase their data."
        elif 'msme' in q_lower:
            answer = "Under MSME Act 2006, buyers must pay MSME suppliers within 45 days. If delayed, compound interest at three times bank rate (approx 21% p.a.) with monthly rests applies."
        elif 'iso' in q_lower or '27001' in q_lower:
            answer = "ISO 27001:2022 has 93 controls across 4 domains. Key requirements include risk assessment, security policy, access control, incident management, and business continuity."
        elif 'employment' in q_lower or 'employee' in q_lower:
            answer = "Employment contracts should include: job title and duties, salary and benefits, probation period (3-6 months), notice period for termination, confidentiality clause, IP ownership, leave policy. Standard employment contracts are generally low risk."
        elif 'hey' in q_lower or 'hi' in q_lower or 'hello' in q_lower:
            answer = "Hey there! 👋 I'm EDIE, your contract and compliance AI assistant. I analyze contracts based on their specific type - Employment, NDA, Software, Service, Vendor, etc. Paste any contract and I'll tell you the risk score and missing clauses!"
        else:
            answer = "I'm here to help! Ask me about DPDP Act (₹250 crore penalties), MSME Act (45-day payment rule), ISO 27001, employment contracts, or any contract clause like Liability caps or Data Protection. Paste a contract in the Analyze tab to see detailed analysis!"
    
    chat_sessions[session_id].append({"role": "user", "content": question})
    chat_sessions[session_id].append({"role": "assistant", "content": answer})
    if len(chat_sessions[session_id]) > 10:
        chat_sessions[session_id] = chat_sessions[session_id][-10:]
    
    return {"success": True, "answer": answer, "session_id": session_id}

@app.post("/ai/ask-analysis")
async def ai_ask_analysis(question: str = Form(...), analysis_data: str = Form(...), session_id: str = Form(None)):
    if not OPENROUTER_API_KEY:
        return {"success": False, "error": "OpenRouter API key not configured."}
    
    try:
        data = json.loads(analysis_data)
        
        if not session_id:
            session_id = f"analysis_{data.get('id', uuid.uuid4())}"
        if session_id not in analysis_sessions:
            analysis_sessions[session_id] = []
        
        answer = call_ai_analysis(question, data, analysis_sessions[session_id])
        
        if not answer:
            q_lower = question.lower()
            risk = data.get('risk_assessment', {})
            clauses = data.get('clause_analysis', {})
            risk_score = risk.get('score', 0)
            risk_level = risk.get('level', 'MEDIUM')
            contract_type = data.get('contract_type', 'general')
            
            if 'data protection' in q_lower or 'dpdp' in q_lower:
                answer = f"The Data Protection clause ensures compliance with India's DPDP Act 2023. Penalties can reach ₹250 crore. Your {contract_type} is {'missing' if 'Data Protection' in clauses.get('missing_list', []) else 'present'} this clause."
            elif 'confidentiality' in q_lower:
                answer = "Confidentiality clauses protect trade secrets and sensitive information. They should define what's confidential, exclude public info, and survive termination."
            elif 'risk score' in q_lower:
                answer = f"The risk score of {risk_score}/100 is calculated based on mandatory clauses present, detected risk patterns, and compliance score. For a {contract_type}, this indicates {risk_level.lower()} risk."
            elif 'compliance' in q_lower:
                comp = data.get('compliance', {})
                answer = f"The compliance score of {comp.get('overall_score', 0)}% checks only laws relevant to {contract_type}s. For example, MSME is checked only for vendor/supplier contracts, DPDP only when data is involved."
            elif 'hey' in q_lower or 'hi' in q_lower or 'hello' in q_lower:
                answer = f"Hey there! 👋 I see this {contract_type} has a {risk_level} risk score of {risk_score}/100. Ask me about any specific clause or risk!"
            else:
                answer = f"The risk score of {risk_score}/100 indicates {risk_level} risk for this {contract_type}, based on {clauses.get('missing', 0)} missing clauses. The recommended action is to {risk.get('decision', 'REVIEW')}."
        
        analysis_sessions[session_id].append({"role": "user", "content": question})
        analysis_sessions[session_id].append({"role": "assistant", "content": answer})
        if len(analysis_sessions[session_id]) > 10:
            analysis_sessions[session_id] = analysis_sessions[session_id][-10:]
        
        return {"success": True, "answer": answer, "session_id": session_id}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/generate-report")
async def generate_report(analysis_data: str = Form(...)):
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx not installed"}
    
    try:
        data = json.loads(analysis_data)
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        title = doc.add_heading('EDIE Enterprise - Contract Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Report ID: {data.get('analysis_id', str(uuid.uuid4()))}")
        doc.add_paragraph(f"Contract Type: {data.get('contract_type', 'General').upper()}")
        doc.add_paragraph("")
        
        risk = data.get('risk_assessment', {})
        doc.add_heading('1. RISK ASSESSMENT', level=1)
        doc.add_paragraph(f"Risk Score: {risk.get('score', 0)}/100")
        doc.add_paragraph(f"Risk Level: {risk.get('level', 'Unknown')}")
        doc.add_paragraph(f"Decision: {risk.get('decision', 'REVIEW')}")
        doc.add_paragraph(f"Action: {risk.get('action', 'Review required')}")
        doc.add_paragraph("")
        
        doc.add_heading('2. EXECUTIVE SUMMARY', level=1)
        doc.add_paragraph(data.get('summary', 'No summary available.'))
        if data.get('ai_explanation'):
            doc.add_paragraph(f"\nAI Insight: {data['ai_explanation']}")
        doc.add_paragraph("")
        
        clauses = data.get('clause_analysis', {})
        doc.add_heading('3. CLAUSE ANALYSIS', level=1)
        doc.add_paragraph(f"Total Clauses Checked: {clauses.get('total', 0)}")
        doc.add_paragraph(f"Mandatory Clauses Present: {clauses.get('mandatory_present', 0)} of {clauses.get('mandatory_total', 0)}")
        doc.add_paragraph("")
        
        doc.add_heading('Present Clauses:', level=2)
        for clause in clauses.get('present_list', [])[:12]:
            doc.add_paragraph(f"✓ {clause}", style='List Bullet')
        
        doc.add_heading('Missing Clauses:', level=2)
        for clause in clauses.get('missing_list', [])[:15]:
            doc.add_paragraph(f"✗ {clause}", style='List Bullet')
        doc.add_paragraph("")
        
        compliance = data.get('compliance', {})
        doc.add_heading('4. COMPLIANCE ASSESSMENT', level=1)
        doc.add_paragraph(f"Overall Compliance Score: {compliance.get('overall_score', 0)}%")
        doc.add_paragraph("(Only laws relevant to this contract type are checked)")
        doc.add_paragraph("")
        for std in compliance.get('standards', []):
            doc.add_paragraph(f"{std.get('standard', '')}: {std.get('score', 0)}% - {std.get('status', 'Unknown')}")
        doc.add_paragraph("")
        
        fin = data.get('financial_analysis', {})
        doc.add_heading('5. FINANCIAL IMPACT ANALYSIS', level=1)
        doc.add_paragraph(f"Contract Value: ₹{fin.get('contract_value', 0):,.0f}")
        doc.add_paragraph(f"Term: {fin.get('term_years', 2)} years")
        doc.add_paragraph(f"Total Revenue: ₹{fin.get('total_revenue', 0):,.0f}")
        doc.add_paragraph(f"Expected Profit: ₹{fin.get('expected_profit', 0):,.0f}")
        doc.add_paragraph(f"Risk Exposure: ₹{fin.get('risk_exposure', 0):,.0f}")
        doc.add_paragraph(f"Risk-Adjusted Profit: ₹{fin.get('risk_adjusted_profit', 0):,.0f}")
        doc.add_paragraph(f"ROI: {fin.get('roi_percentage', 0)}%")
        doc.add_paragraph("")
        
        doc.add_heading('6. RECOMMENDATIONS', level=1)
        for rec in data.get('recommendations', [])[:8]:
            doc.add_paragraph(f"• {rec}", style='List Bullet')
        doc.add_paragraph("")
        
        doc.add_heading('7. FINAL DECISION', level=1)
        doc.add_paragraph(f"Decision: {risk.get('decision', 'REVIEW')}")
        doc.add_paragraph(f"Action Required: {risk.get('action', 'Review required')}")
        
        os.makedirs("reports", exist_ok=True)
        report_filename = f"EDIE_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = os.path.join("reports", report_filename)
        doc.save(report_path)
        
        return {"success": True, "report_path": report_path, "filename": report_filename}
    except Exception as e:
        logger.error(f"Report generation error: {e}")
        return {"success": False, "error": str(e)}

@app.get("/download-report/{filename}")
async def download_report(filename: str):
    filepath = os.path.join("reports", filename)
    if os.path.exists(filepath):
        return FileResponse(filepath, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)
    return {"error": "Report not found"}

@app.get("/history")
async def get_history():
    return {"success": True, "history": storage.get_all(20)}

@app.get("/knowledge-base/status")
async def get_kb_status():
    return {"files_loaded": len(kb.loaded_files), "status": "ready"}

@app.post("/clear-history")
async def clear_history():
    storage.analyses.clear()
    return {"success": True}

if __name__ == "__main__":
    import uvicorn
    
    print("\n" + "="*70)
    print("🚀 EDIE ENTERPRISE v5.0 - Contract-Type Specific Analysis")
    print("="*70)
    print(f"📚 Knowledge Base: {len(kb.loaded_files)} files loaded")
    print(f"🤖 AI: {'✅ Enabled (OpenRouter/Gemma)' if OPENROUTER_API_KEY else '❌ Disabled'}")
    print(f"📄 Contract Types Supported: Employment, NDA, Software, Service, Consulting, Vendor, Purchase, Lease, Maintenance, Distribution, Joint Venture, General")
    print("="*70)
    print("🌐 Server: http://127.0.0.1:8000")
    print("="*70)
    
    if not OPENROUTER_API_KEY:
        print("\n⚠️ To enable AI, set environment variable:")
        print("   PowerShell: $env:OPENROUTER_API_KEY = 'sk-or-v1-your-key'")
        print("   Get free key: https://openrouter.ai/keys")
        print("="*70)
    
    uvicorn.run(app, host="0.0.0.0", port=8000)